
from openai import OpenAI
import os
from flask import Flask, request, send_file, make_response
from flask_cors import CORS
from datetime import datetime
import io
from docx import Document

client = OpenAI(api_key='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')

app = Flask(__name__)
CORS(app, origins=["chrome-extension://gbbfcbcjpdlabjfeccljliaedcpfnnpg"])

@app.route("/customize", methods=["POST"])
def customize():
    data = request.get_json()
    job_desc = data.get("jobDesc", "")

    prompt = f"""You are a resume optimization assistant. Tailor the following resume bullets to match the job description below.
- Do not change the core content, but rewrite the bullets to match the tone, technologies, and priorities of the job.
- Include quantifiable metrics if available.
- Return exactly 10 bullets: 2 for UNIVERSITY OF ILLINOIS, 3 for EXTUENT, 5 for FRAPPE.
- Do not use company names in the bullet openings.

Return this JSON exactly:

{{
  "experience": [
    "• Bullet 1",
    "• Bullet 2",
    "• Bullet 3",
    "• Bullet 4",
    "• Bullet 5",
    "• Bullet 6",
    "• Bullet 7",
    "• Bullet 8",
    "• Bullet 9",
    "• Bullet 10"
  ],
  "skills": "Skill1 | Skill2 | Skill3 | ..."
}}

Only return valid JSON. Do not explain, apologize, or use markdown.

Base EXPERIENCE:
UNIVERSITY OF ILLINOIS URBANA-CHAMPAIGN
Product Data Analyst - Research Assistant (Mar 2024 – Aug 2024)
- Launching an AI-based product with OSF Healthcare and University of Illinois to optimize rural healthcare allocation.
- Managed extensive zip-code level open-access healthcare data using G-Suite to enhance service accessibility.

EXTUENT
Product Manager (Mar 2023 – Aug 2023)
- Drove UX improvement across ERP lifecycle using ERPNext, boosting productivity by 15% and reducing process time by 20%.
- Led 5 enterprise implementations with KPI-focused roadmaps, increasing client satisfaction by 15% and loyalty rates.
- Reduced service disruptions by 25% and downtime by 30% through stakeholder engagement and issue root cause discovery.

FRAPPE
Project Manager (Oct 2020 – Jan 2023)
- Delivered 30+ ERP projects 20% ahead of schedule across retail and manufacturing, optimizing order processing systems.
- Increased project efficiency by 20% through Agile and JIRA, and achieved 95% on-time delivery rate.
- Boosted client satisfaction by 20% via data-driven decisions and Qualtrics-based feedback collection.
- Reduced support tickets by 25% through clear documentation, blogs, and help articles on Notion and ERPNext.
- Executed API integrations and cloud migration across 4 full-lifecycle projects using Waterfall methodology.

JOB DESCRIPTION:
{job_desc}
"""

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            { "role": "user", "content": prompt }
        ]
    )

    parsed = eval(response.choices[0].message.content)

    doc = Document("base_resume.docx")
    experience = parsed["experience"]
    skills = parsed["skills"]

    def replace_paragraphs(section_title, new_bullets, n):
        for i in range(len(doc.paragraphs)-1, -1, -1):
            if section_title in doc.paragraphs[i].text:
                # Remove next n paragraphs
                for _ in range(n):
                    if i + 1 < len(doc.paragraphs):
                        del doc.paragraphs[i + 1]
                for j in range(n):
                    doc.paragraphs.insert(i + 1 + j, doc.add_paragraph(new_bullets[j]))
                break

    replace_paragraphs("UNIVERSITY OF ILLINOIS URBANA-CHAMPAIGN", experience[0:2], 2)
    replace_paragraphs("EXTUENT", experience[2:5], 3)
    replace_paragraphs("FRAPPE", experience[5:10], 5)

    for para in doc.paragraphs:
        if "Core Competencies" in para.text:
            para.text = "Core Competencies - " + skills
            break

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = "Vrinda Menon Resume - " + datetime.now().strftime("%Y-%m-%d") + ".docx"
    response = make_response(send_file(output, as_attachment=True, download_name=filename))
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)
